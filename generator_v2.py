import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Mm  # регулирование полей

class SubtaskFrame(tk.Frame):
    def __init__(self, parent, subtask_id, delete_callback, app):
        super().__init__(parent, relief=tk.GROOVE, borderwidth=2)
        self.pack(fill=tk.X, padx=10, pady=5)
        
        self.subtask_id = subtask_id
        self.delete_callback = delete_callback
        self.app = app  # Сохраняем ссылку на основное приложение
        self.images = []  # Список путей к изображениям
        self.tables = []  # Список таблиц (словари с размерами и данными)
        
        # Заголовок подпункта
        header_frame = tk.Frame(self)
        header_frame.pack(fill=tk.X)
        
        self.label = tk.Label(header_frame, text=f"Подпункт {self.subtask_id}")
        self.label.pack(side=tk.LEFT, padx=5)
        
        delete_btn = tk.Button(header_frame, text="Удалить", command=self.delete_self)
        delete_btn.pack(side=tk.RIGHT, padx=5)
        
        # Текст подпункта
        tk.Label(self, text="Текст:").pack(anchor=tk.W, padx=5)
        self.text = tk.Text(self, height=4, width=70)
        self.text.pack(padx=5, pady=2, fill=tk.X)
        
        # Баллы за подпункт
        points_frame = tk.Frame(self)
        points_frame.pack(fill=tk.X, padx=5)
        tk.Label(points_frame, text="Баллы:").pack(side=tk.LEFT)
        self.points_entry = tk.Entry(points_frame, width=5)
        self.points_entry.pack(side=tk.LEFT, padx=5)
        
        # Кнопки для изображений и таблиц
        buttons_frame = tk.Frame(self)
        buttons_frame.pack(fill=tk.X, padx=5, pady=5)
        
        add_image_btn = tk.Button(buttons_frame, text="Добавить изображение", command=self.add_image)
        add_image_btn.pack(side=tk.LEFT, padx=5)
        
        add_table_btn = tk.Button(buttons_frame, text="Добавить таблицу", command=self.add_table)
        add_table_btn.pack(side=tk.LEFT, padx=5)
        
        # Информация о добавленных элементах
        self.elements_info = tk.Label(self, text="Элементы: нет")
        self.elements_info.pack(anchor=tk.W, padx=5)

        self.app.setup_text_bindings(self.text)
        self.app.setup_text_bindings(self.points_entry)

        answer_box_frame = tk.Frame(self)
        answer_box_frame.pack(fill=tk.X, padx=5)
        tk.Label(answer_box_frame, text="Высота места для ответа (см):").pack(side=tk.LEFT)
        self.answer_height_var = tk.StringVar(value="3")
        tk.Spinbox(answer_box_frame, from_=1, to=15, textvariable=self.answer_height_var, width=5).pack(side=tk.LEFT, padx=5)
        
    def delete_self(self):
        self.delete_callback(self.subtask_id)
        self.destroy()
        
    def add_image(self):
        filename = filedialog.askopenfilename(filetypes=[("Изображения", "*.png *.jpg *.jpeg *.bmp *.emf")])
        if filename:
            self.images.append(filename)
            self.update_elements_info()
    
    def add_table(self):
        dialog = TableDialog(self)
        if dialog.result:
            self.tables.append(dialog.result)
            self.update_elements_info()
    
    def update_elements_info(self):
        info = f"Элементы: {len(self.images)} изображений, {len(self.tables)} таблиц"

    def get_data(self):
        return {
            "text": self.text.get("1.0", tk.END).strip(),
            "points": self.points_entry.get().strip(),
            "images": self.images.copy(),
            "tables": self.tables.copy(),
            "answer_height": float(self.answer_height_var.get()) if self.answer_height_var.get().strip() else 3.0
        }


class TableDialog(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Создание таблицы")
        self.geometry("400x300")
        self.result = None
        
        # Размеры таблицы
        size_frame = tk.Frame(self)
        size_frame.pack(pady=10)
        
        tk.Label(size_frame, text="Строки:").grid(row=0, column=0, padx=5)
        self.rows_var = tk.IntVar(value=2)
        tk.Spinbox(size_frame, from_=1, to=10, textvariable=self.rows_var, width=5).grid(row=0, column=1)
        
        tk.Label(size_frame, text="Столбцы:").grid(row=0, column=2, padx=5)
        self.cols_var = tk.IntVar(value=2)
        tk.Spinbox(size_frame, from_=1, to=10, textvariable=self.cols_var, width=5).grid(row=0, column=3)
        
        # Кнопка создания ячеек для заполнения
        tk.Button(self, text="Создать ячейки", command=self.create_cells).pack(pady=5)
        
        # Фрейм для ячеек (будет создан динамически)
        self.cells_frame = tk.Frame(self)
        self.cells_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Кнопки Ок/Отмена
        buttons_frame = tk.Frame(self)
        buttons_frame.pack(side=tk.BOTTOM, pady=10)
        
        tk.Button(buttons_frame, text="ОК", command=self.ok_command).pack(side=tk.LEFT, padx=10)
        tk.Button(buttons_frame, text="Отмена", command=self.cancel_command).pack(side=tk.LEFT)
        
        self.cell_entries = []
        self.transient(parent)
        self.grab_set()
        parent.wait_window(self)
    
    def create_cells(self):
        # Очистить предыдущие ячейки
        for widget in self.cells_frame.winfo_children():
            widget.destroy()
        
        self.cell_entries = []
        rows = self.rows_var.get()
        cols = self.cols_var.get()
        
        # Создать заголовки столбцов
        for col in range(cols):
            tk.Label(self.cells_frame, text=f"Стлб {col+1}").grid(row=0, column=col+1, padx=2, pady=2)
        
        # Создать ячейки для ввода
        for row in range(rows):
            tk.Label(self.cells_frame, text=f"Стр {row+1}").grid(row=row+1, column=0, padx=2, pady=2)
            row_entries = []
            
            for col in range(cols):
                entry = tk.Entry(self.cells_frame, width=10)
                entry.grid(row=row+1, column=col+1, padx=2, pady=2)
                row_entries.append(entry)
                
            self.cell_entries.append(row_entries)
    
    def ok_command(self):
        if not self.cell_entries:
            messagebox.showwarning("Внимание", "Сначала создайте ячейки таблицы")
            return
            
        rows = self.rows_var.get()
        cols = self.cols_var.get()
        
        # Собрать данные из ячеек
        table_data = []
        for row in range(rows):
            row_data = []
            for col in range(cols):
                cell_value = self.cell_entries[row][col].get()
                row_data.append(cell_value)
            table_data.append(row_data)
        
        self.result = {
            "rows": rows,
            "cols": cols,
            "data": table_data
        }
        
        self.destroy()
    
    def cancel_command(self):
        self.destroy()


class ChemistryTaskGenerator:
    def __init__(self, root):
        self.root = root
        self.tasks = []
        self.image_path_var = tk.StringVar()
        self.task_type_var = tk.StringVar(value="Органическая химия")
        self.subtask_frames = {}
        self.next_subtask_id = 1
        
        self.create_ui()
    def setup_text_bindings(self, text_widget):
        """Настраивает биндинги для текстового виджета для корректной работы копирования/вставки"""
        # Для Windows/Linux
        text_widget.bind("<Control-v>", lambda event: self.paste_text(event, text_widget))
        text_widget.bind("<Control-c>", lambda event: self.copy_text(event, text_widget))
        text_widget.bind("<Control-x>", lambda event: self.cut_text(event, text_widget))
        
        # Для macOS
        text_widget.bind("<Command-v>", lambda event: self.paste_text(event, text_widget))
        text_widget.bind("<Command-c>", lambda event: self.copy_text(event, text_widget))
        text_widget.bind("<Command-x>", lambda event: self.cut_text(event, text_widget))
        
        # Добавляем контекстное меню
        self.create_context_menu(text_widget)

    def paste_text(self, event, text_widget):
        """Вставка текста из буфера обмена"""
        try:
            text = self.root.clipboard_get()
            if event.widget.cget("state") != "disabled":
                if isinstance(event.widget, tk.Text):
                    event.widget.insert(tk.INSERT, text)
                elif isinstance(event.widget, tk.Entry):
                    event.widget.insert(tk.INSERT, text)
            return "break"  # Предотвращает обработку события дефолтным обработчиком
        except:
            return

    def copy_text(self, event, text_widget):
        """Копирование текста в буфер обмена"""
        try:
            if isinstance(event.widget, tk.Text):
                if not event.widget.tag_ranges(tk.SEL):
                    return
                text = event.widget.get(tk.SEL_FIRST, tk.SEL_LAST)
            elif isinstance(event.widget, tk.Entry):
                if not event.widget.selection_present():
                    return
                text = event.widget.selection_get()
            else:
                return
            self.root.clipboard_clear()
            self.root.clipboard_append(text)
        except:
            pass
        return "break"

    def cut_text(self, event, text_widget):
        """Вырезание текста в буфер обмена"""
        if self.copy_text(event, text_widget) != "break":
            return
        try:
            if isinstance(event.widget, tk.Text):
                event.widget.delete(tk.SEL_FIRST, tk.SEL_LAST)
            elif isinstance(event.widget, tk.Entry):
                event.widget.delete(tk.SEL_FIRST, tk.SEL_LAST)
        except:
            pass
        return "break"

    def create_context_menu(self, widget):
        """Создает контекстное меню для текстового виджета"""
        context_menu = tk.Menu(widget, tearoff=0)
        context_menu.add_command(label="Вырезать", command=lambda: self.cut_text_menu(widget))
        context_menu.add_command(label="Копировать", command=lambda: self.copy_text_menu(widget))
        context_menu.add_command(label="Вставить", command=lambda: self.paste_text_menu(widget))
        
        # Привязываем меню к правой кнопке мыши
        widget.bind("<Button-3>", lambda event: self.show_context_menu(event, context_menu))

    def cut_text_menu(self, widget):
        """Вырезание текста через меню"""
        self.copy_text_menu(widget)
        try:
            if isinstance(widget, tk.Text):
                widget.delete(tk.SEL_FIRST, tk.SEL_LAST)
            elif isinstance(widget, tk.Entry):
                widget.delete(tk.SEL_FIRST, tk.SEL_LAST)
        except:
            pass

    def copy_text_menu(self, widget):
        """Копирование текста через меню"""
        try:
            if isinstance(widget, tk.Text):
                if not widget.tag_ranges(tk.SEL):
                    return
                text = widget.get(tk.SEL_FIRST, tk.SEL_LAST)
            elif isinstance(widget, tk.Entry):
                if not widget.selection_present():
                    return
                text = widget.selection_get()
            else:
                return
            self.root.clipboard_clear()
            self.root.clipboard_append(text)
        except:
            pass

    def paste_text_menu(self, widget):
        """Вставка текста через меню"""
        try:
            text = self.root.clipboard_get()
            if isinstance(widget, tk.Text):
                widget.insert(tk.INSERT, text)
            elif isinstance(widget, tk.Entry):
                widget.insert(tk.INSERT, text)
        except:
            pass

    def show_context_menu(self, event, menu):
        """Показывает контекстное меню в месте клика"""
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()
    
    def create_ui(self):
        # Создаем notebook для вкладок
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Вкладка для создания задач
        self.task_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.task_frame, text="Создание задачи")
        
        # Вкладка для просмотра всех задач
        self.tasks_list_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.tasks_list_frame, text="Список задач")
        
        # Компоненты для создания задач
        self.create_task_ui()
        
        # Компоненты для просмотра задач
        self.create_tasks_list_ui()
    
    def create_task_ui(self):
        # Основная информация о задаче
        info_frame = ttk.LabelFrame(self.task_frame, text="Основная информация")
        info_frame.pack(fill=tk.X, padx=10, pady=5)
        
        # Тип задачи
        type_frame = tk.Frame(info_frame)
        type_frame.pack(fill=tk.X, pady=2)
        
        tk.Label(type_frame, text="Тип задачи:").pack(side=tk.LEFT, padx=5)
        
        task_type_menu = ttk.OptionMenu(
            type_frame, 
            self.task_type_var, 
            "Органическая химия",  # Значение по умолчанию
            "Органическая химия",  # Добавлено как опция
            "Неорганическая химия", 
            "Физическая химия",
            "Аналитическая химия"
        )
        task_type_menu.pack(side=tk.LEFT, padx=5)
        
        # Название задачи и баллы
        title_frame = tk.Frame(info_frame)
        title_frame.pack(fill=tk.X, pady=2)
        
        tk.Label(title_frame, text="Название задачи:").pack(side=tk.LEFT, padx=5)
        self.title_entry = tk.Entry(title_frame, width=40)
        self.title_entry.pack(side=tk.LEFT, padx=5)
        
        tk.Label(title_frame, text="Баллы:").pack(side=tk.LEFT, padx=5)
        self.points_entry = tk.Entry(title_frame, width=5)
        self.points_entry.pack(side=tk.LEFT, padx=5)
        
        # Условие задачи
        statement_frame = ttk.LabelFrame(self.task_frame, text="Условие задачи")
        statement_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.statement_text = tk.Text(statement_frame, height=8, width=80)
        self.statement_text.pack(padx=5, pady=5, fill=tk.BOTH, expand=True)
        
        # Изображение к основному условию
        image_frame = tk.Frame(statement_frame)
        image_frame.pack(fill=tk.X, pady=2)
        
        tk.Label(image_frame, text="Изображение:").pack(side=tk.LEFT, padx=5)
        tk.Entry(image_frame, textvariable=self.image_path_var, width=40).pack(side=tk.LEFT, padx=5)
        tk.Button(image_frame, text="Выбрать файл...", command=self.choose_image).pack(side=tk.LEFT, padx=5)
        
        # Подпункты задачи
        subtasks_frame = ttk.LabelFrame(self.task_frame, text="Подпункты задачи")
        subtasks_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # Кнопка добавления подпункта
        tk.Button(subtasks_frame, text="Добавить подпункт", command=self.add_subtask).pack(anchor=tk.W, padx=10, pady=5)
        
        # Контейнер для подпунктов (с прокруткой)
        subtasks_container = tk.Frame(subtasks_frame)
        subtasks_container.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        self.subtasks_canvas = tk.Canvas(subtasks_container)
        scrollbar = tk.Scrollbar(subtasks_container, orient=tk.VERTICAL, command=self.subtasks_canvas.yview)
        self.subtasks_frame = tk.Frame(self.subtasks_canvas)
        
        self.subtasks_canvas.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.subtasks_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.subtasks_canvas.create_window((0, 0), window=self.subtasks_frame, anchor=tk.NW)
        self.subtasks_frame.bind("<Configure>", lambda e: self.subtasks_canvas.configure(scrollregion=self.subtasks_canvas.bbox("all")))
        
        # Кнопки действий
        actions_frame = tk.Frame(self.task_frame)
        actions_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Button(actions_frame, text="Добавить задачу", command=self.add_task).pack(side=tk.LEFT, padx=5)
        tk.Button(actions_frame, text="Очистить форму", command=self.clear_form).pack(side=tk.LEFT, padx=5)
        tk.Button(actions_frame, text="Создать Word-документ", command=self.export_docx).pack(side=tk.RIGHT, padx=5)

        self.setup_text_bindings(self.title_entry)
        self.setup_text_bindings(self.statement_text)
        self.setup_text_bindings(self.points_entry)


    def create_tasks_list_ui(self):
        # Список всех добавленных задач
        tk.Label(self.tasks_list_frame, text="Добавленные задачи:").pack(anchor=tk.W, padx=10, pady=5)
        
        # Фрейм для списка с прокруткой
        list_container = tk.Frame(self.tasks_list_frame)
        list_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.tasks_listbox = tk.Listbox(list_container, height=15, width=70)
        scrollbar = tk.Scrollbar(list_container, orient=tk.VERTICAL, command=self.tasks_listbox.yview)
        
        self.tasks_listbox.configure(yscrollcommand=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.tasks_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Кнопки управления списком
        buttons_frame = tk.Frame(self.tasks_list_frame)
        buttons_frame.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Button(buttons_frame, text="Удалить выбранную задачу", command=self.delete_selected_task).pack(side=tk.LEFT, padx=5)
        tk.Button(buttons_frame, text="Создать Word-документ", command=self.export_docx).pack(side=tk.RIGHT, padx=5)
    

    def add_subtask(self):
        subtask_frame = SubtaskFrame(self.subtasks_frame, self.next_subtask_id, self.delete_subtask, self)
        self.subtask_frames[self.next_subtask_id] = subtask_frame
        self.next_subtask_id += 1
    
    def delete_subtask(self, subtask_id):
        if subtask_id in self.subtask_frames:
            del self.subtask_frames[subtask_id]
    
    def choose_image(self):
        filename = filedialog.askopenfilename(filetypes=[("Изображения", "*.png *.jpg *.jpeg *.bmp *.emf")])
        if filename:
            self.image_path_var.set(filename)
    
    def add_task(self):
        task_type = self.task_type_var.get()
        title = self.title_entry.get().strip()
        statement = self.statement_text.get("1.0", tk.END).strip()
        points = self.points_entry.get().strip()
        image_path = self.image_path_var.get()

        if not title or not statement or not points:
            messagebox.showerror("Ошибка", "Название, условие и баллы обязательны для заполнения.")
            return

        # Собираем подпункты
        subtasks = []
        for subtask_id, frame in self.subtask_frames.items():
            subtask_data = frame.get_data()
            if subtask_data["text"]:  # Добавляем только подпункты с текстом
                subtasks.append(subtask_data)

        self.tasks.append({
            "type": task_type,
            "title": title,
            "statement": statement,
            "points": points,
            "image": image_path,
            "subtasks": subtasks
        })

        # Обновить список задач
        self.update_tasks_list()
        
        self.clear_form()
        messagebox.showinfo("Добавлено", f"Задача '{title}' добавлена.")
    
    def update_tasks_list(self):
        self.tasks_listbox.delete(0, tk.END)
        for i, task in enumerate(self.tasks, start=1):
            self.tasks_listbox.insert(tk.END, f"{i}. {task['title']} ({task['points']} баллов)")
    
    def delete_selected_task(self):
        selected = self.tasks_listbox.curselection()
        if not selected:
            messagebox.showwarning("Внимание", "Выберите задачу для удаления")
            return
            
        index = selected[0]
        task_title = self.tasks[index]['title']
        del self.tasks[index]
        self.update_tasks_list()
        messagebox.showinfo("Удалено", f"Задача '{task_title}' удалена")
    
    def clear_form(self):
        self.title_entry.delete(0, tk.END)
        self.statement_text.delete("1.0", tk.END)
        self.points_entry.delete(0, tk.END)
        self.image_path_var.set("")
        
        # Удалить все подпункты
        for frame in list(self.subtask_frames.values()):
            frame.destroy()
        self.subtask_frames.clear()
        self.next_subtask_id = 1
    
    def create_table(self, doc, table_data):
        rows = table_data["rows"]
        cols = table_data["cols"]
        data = table_data["data"]
        
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # Заполнить таблицу данными
        for i in range(rows):
            for j in range(cols):
                cell = table.cell(i, j)
                cell.text = data[i][j]
                # Центрирование текста в ячейках
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        return table
    
    def add_answer_box(self, doc, height_cm=3):
        """Создаёт рамку для ответа указанной высоты (в см)"""
        # Добавляем пустой параграф для рамки
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Создаем элемент "рамка"
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        
        # Добавляем границы со всех сторон
        borders = ['top', 'left', 'bottom', 'right']
        for border in borders:
            bdr = OxmlElement(f'w:{border}')
            bdr.set(qn('w:val'), 'single')
            bdr.set(qn('w:sz'), '24')  # толщина границы (в восьмых пункта)
            bdr.set(qn('w:space'), '0')
            bdr.set(qn('w:color'), '000000')
            pBdr.append(bdr)
        
        pPr.append(pBdr)
        
        # Устанавливаем высоту абзаца
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        # Конвертируем см в twips (единица измерения Word - 1/20 пункта)
        height_twips = int(height_cm * 567)  # 1 см ≈ 567 twips
        spacing.set(qn('w:line'), str(height_twips))
        spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(spacing)
        
        return p

    def set_paragraph_justify(self, paragraph):
        """Устанавливает выравнивание по ширине для параграфа"""
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
    def export_docx(self):
        if not self.tasks:
            messagebox.showerror("Ошибка", "Нет задач для экспорта.")
            return

        doc = Document()
        
        # Настройка стилей документа
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        # setting up margins
        for section in doc.sections:
            section.left_margin = Cm(1.5)    # 1.25 дюйма (3.17 см)
            section.right_margin = Cm(1.5)   # 1.25 дюйма
            section.top_margin = Cm(1.5)        # 1 дюйм (2.54 см)
            section.bottom_margin = Cm(1.5)     # 1 дюйм
            section.header_distance = Cm(0.5) # 0.5 дюйма до верхнего колонтитула
            section.footer_distance = Cm(0.5)
        # Заголовок документа
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("Олимпиада НАШ СЛОН, сезон 1\n Апрель 2025")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Проверяем, что у нас есть задачи для добавления
        if self.tasks:  # Проверка, что список tasks не пустой
            # Добавление задач
            for i, task in enumerate(self.tasks, start=1):
                # Заголовок задачи
                task_heading = doc.add_paragraph(f"Задача {i}. {task['title']} ({task['points']} баллов)", style='Heading 1')
                task_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Условие задачи
                doc.add_paragraph("Условие:", style='Heading 2')
                for line in task['statement'].split('\n'):
                    para = doc.add_paragraph(line.strip())
                    self.set_paragraph_justify(para)  # Выравнивание по ширине
                
                # Изображение к основному условию, если есть
                if task.get("image"):  # Используем .get() для безопасного доступа к ключу
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p.add_run().add_picture(task["image"], width=Inches(4))
                    except Exception as e:
                        doc.add_paragraph(f"[Не удалось вставить изображение: {str(e)}]")
                
                # Добавление подпунктов
                if task.get("subtasks"):  # Используем .get() для безопасного доступа к ключу
                    for j, subtask in enumerate(task["subtasks"], start=1):
                        subtask_para = doc.add_paragraph(f"{j}) {subtask['text']}")
                        self.set_paragraph_justify(subtask_para)  # Выравнивание по ширине
                        
                        if subtask.get("points"):  # Используем .get() для безопасного доступа к ключу
                            subtask_para.add_run(f" ({subtask['points']} баллов)")
                        
                        # Добавление изображений к подпункту
                        for img_path in subtask.get("images", []):  # Используем .get() с пустым списком по умолчанию
                            try:
                                p = doc.add_paragraph()
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                                p.add_run().add_picture(img_path, width=Inches(3.5))
                            except Exception as e:
                                doc.add_paragraph(f"[Не удалось вставить изображение: {str(e)}]")
                        
                        # Добавление таблиц к подпункту
                        for table_data in subtask.get("tables", []):  # Используем .get() с пустым списком по умолчанию
                            doc.add_paragraph()
                            self.create_table(doc, table_data)
                            doc.add_paragraph()
                        
                        # Место для ответа на подпункт - рамка
                        # Размер рамки может зависеть от типа задачи или настроек подпункта
                        height = subtask.get("answer_height", 3.0)  # Используем .get() для безопасного доступа
                        self.add_answer_box(doc, height_cm=height)
                        
                        doc.add_paragraph()
                else:
                    # Место для ответа в зависимости от типа задачи
                    doc.add_paragraph("Ответ:", style='Heading 2')
                    
                    task_type = task.get("type", "")  # Используем .get() для безопасного доступа
                    
                    if task_type in ["Органическая химия", "Неорганическая химия"]:
                        # Таблица для структурных формул
                        table = doc.add_table(rows=2, cols=2)
                        table.style = 'Table Grid'
                        for row in table.rows:
                            for cell in row.cells:
                                cell.text = " "
                                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    elif task_type == "Физическая химия":
                        # Место для расчетов - большая рамка
                        self.add_answer_box(doc, height_cm=8)
                        
                        p = doc.add_paragraph("Ответ: ")
                        self.add_answer_box(doc, height_cm=1.5)  # меньшая рамка для финального ответа
                    
                    else:
                        # Для прочих типов задач
                        self.add_answer_box(doc, height_cm=4)
                
                # Разделитель между задачами
                doc.add_page_break()

        # Диалог сохранения файла
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word файлы", "*.docx")],
            title="Сохранить как"
        )
        
        if save_path:
            try:
                doc.save(save_path)
                messagebox.showinfo("Готово", f"Файл сохранён: {save_path}")
            except Exception as e:
                messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить файл: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    root.title("Генератор задач для химической олимпиады")
    root.geometry("800x700")
    
    app = ChemistryTaskGenerator(root)
    
    root.mainloop()